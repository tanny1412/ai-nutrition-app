from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple
from uuid import uuid4

from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
try:
    from langchain_community.vectorstores.pgvector import PGVector
except ImportError:  # pragma: no cover - optional dependency
    PGVector = None  # type: ignore
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from rank_bm25 import BM25Okapi
import tiktoken

from ..config import get_settings
from .memory_service import MemoryService, get_memory_service


@dataclass
class DocumentChunk:
    """Lightweight representation of a chunk with metadata."""

    content: str
    metadata: Dict[str, object]


class VectorStoreProvider:
    """Abstract interface for vector store implementations."""

    def add_documents(self, chunks: Sequence[DocumentChunk]) -> None:  # pragma: no cover - interface only
        raise NotImplementedError

    def similarity_search(self, query: str, k: int) -> List[Dict[str, object]]:  # pragma: no cover - interface only
        raise NotImplementedError

    def list_documents(self) -> List[DocumentChunk]:  # pragma: no cover - interface only
        raise NotImplementedError

    def persist(self) -> None:  # pragma: no cover - interface only
        raise NotImplementedError

    def is_empty(self) -> bool:  # pragma: no cover - interface only
        raise NotImplementedError


class FaissVectorStoreProvider(VectorStoreProvider):
    """Default FAISS-backed vector store with on-disk persistence."""

    def __init__(self, store_path: Path, embeddings: OpenAIEmbeddings) -> None:
        self._store_path = store_path
        self._embeddings = embeddings
        self._store: Optional[FAISS] = None
        self._load()

    def _load(self) -> None:
        if self._store_path.exists():
            self._store = FAISS.load_local(
                str(self._store_path),
                self._embeddings,
                allow_dangerous_deserialization=True,
            )

    def _ensure_store(self, documents: Sequence[Document]) -> None:
        if self._store is None:
            if not documents:
                return
            self._store = FAISS.from_documents(list(documents), self._embeddings)
        else:
            self._store.add_documents(list(documents))

    def add_documents(self, chunks: Sequence[DocumentChunk]) -> None:
        documents = [Document(page_content=chunk.content, metadata=chunk.metadata) for chunk in chunks]
        self._ensure_store(documents)
        self.persist()

    def similarity_search(self, query: str, k: int) -> List[Dict[str, object]]:
        if self._store is None:
            return []
        results = self._store.similarity_search_with_score(query, k=k)
        candidates: List[Dict[str, object]] = []
        for doc, score in results:
            metadata = dict(doc.metadata or {})
            similarity = 1.0 / (1.0 + float(score)) if score is not None else 0.0
            candidates.append(
                {
                    "content": doc.page_content,
                    "metadata": metadata,
                    "dense_score": similarity,
                },
            )
        return candidates

    def list_documents(self) -> List[DocumentChunk]:
        if self._store is None:
            return []
        chunks: List[DocumentChunk] = []
        for doc in self._store.docstore._dict.values():  # pylint: disable=protected-access
            chunks.append(DocumentChunk(content=doc.page_content, metadata=dict(doc.metadata or {})))
        return chunks

    def persist(self) -> None:
        if self._store is not None:
            self._store.save_local(str(self._store_path))

    def is_empty(self) -> bool:
        if self._store is None:
            return True
        return len(self._store.docstore._dict) == 0  # pylint: disable=protected-access


class PgVectorStoreProvider(VectorStoreProvider):
    """Managed Postgres/pgvector backend."""

    def __init__(self, dsn: str, embeddings: OpenAIEmbeddings, collection_name: str = "nutrition_rag") -> None:
        if not dsn:
            raise RuntimeError(
                "VECTOR_BACKEND is set to pgvector but POSTGRES_DSN is empty. Provide a DSN or switch back to FAISS.",
            )
        if PGVector is None:  # pragma: no cover - optional dependency
            raise RuntimeError("pgvector backend requires langchain-community PGVector support to be installed.")
        self._dsn = dsn
        self._embeddings = embeddings
        self._collection_name = collection_name
        self._store = PGVector(
            connection_string=dsn,
            collection_name=collection_name,
            embedding_function=embeddings,
            use_jsonb=True,
        )

    def add_documents(self, chunks: Sequence[DocumentChunk]) -> None:
        documents = [Document(page_content=chunk.content, metadata=chunk.metadata) for chunk in chunks]
        if not documents:
            return
        self._store.add_documents(documents)

    def similarity_search(self, query: str, k: int) -> List[Dict[str, object]]:
        if k <= 0:
            return []
        results = self._store.similarity_search_with_score(query, k=k)
        candidates: List[Dict[str, object]] = []
        for doc, score in results:
            metadata = dict(doc.metadata or {})
            similarity = 1.0 / (1.0 + float(score)) if score is not None else 0.0
            candidates.append(
                {
                    "content": doc.page_content,
                    "metadata": metadata,
                    "dense_score": similarity,
                },
            )
        return candidates

    def list_documents(self) -> List[DocumentChunk]:
        # PGVector does not expose a direct listing API; fall back to dense search over a large window.
        # If the store is empty this returns [] quickly.
        documents: List[DocumentChunk] = []
        # Attempt to fetch up to 200 documents using similarity search with an empty query.
        try:
            results = self._store.similarity_search("", k=200)
        except Exception:  # pragma: no cover - best effort only
            results = []
        for doc in results:
            documents.append(DocumentChunk(content=doc.page_content, metadata=dict(doc.metadata or {})))
        return documents

    def persist(self) -> None:  # pragma: no cover - pgvector persistence handled by database
        return

    def is_empty(self) -> bool:
        try:
            count = self._store._collection.count()  # pylint: disable=protected-access
        except Exception:  # pragma: no cover - best effort only
            return False
        return count == 0


class BM25KeywordIndex:
    """Keyword retrieval using BM25."""

    def __init__(self) -> None:
        self._documents: List[str] = []
        self._metadata: List[Dict[str, object]] = []
        self._tokenized: List[List[str]] = []
        self._bm25: Optional[BM25Okapi] = None

    @staticmethod
    def _tokenize(text: str) -> List[str]:
        return re.findall(r"\w+", text.lower())

    def add_chunks(self, chunks: Sequence[DocumentChunk]) -> None:
        for chunk in chunks:
            self._documents.append(chunk.content)
            self._metadata.append(chunk.metadata)
            self._tokenized.append(self._tokenize(chunk.content))
        if self._documents:
            self._bm25 = BM25Okapi(self._tokenized)

    def search(self, query: str, k: int) -> List[Dict[str, object]]:
        if not self._documents or self._bm25 is None:
            return []
        tokens = self._tokenize(query)
        scores = self._bm25.get_scores(tokens)
        ranked = sorted(enumerate(scores), key=lambda item: item[1], reverse=True)[:k]
        results: List[Dict[str, object]] = []
        for idx, score in ranked:
            results.append(
                {
                    "content": self._documents[idx],
                    "metadata": self._metadata[idx],
                    "bm25_score": float(score),
                },
            )
        return results


class BaseReranker:
    def rerank(
        self,
        query: str,
        candidates: Sequence[Dict[str, object]],
        top_n: int,
    ) -> List[Dict[str, object]]:  # pragma: no cover - interface only
        raise NotImplementedError


class CohereReranker(BaseReranker):
    def __init__(self, api_key: str, model: str = "rerank-3.5") -> None:
        import cohere

        if not api_key:
            raise RuntimeError("COHERE_API_KEY must be provided for the Cohere reranker.")
        self._client = cohere.Client(api_key)
        self._model = model

    def rerank(
        self,
        query: str,
        candidates: Sequence[Dict[str, object]],
        top_n: int,
    ) -> List[Dict[str, object]]:
        if not candidates:
            return []
        documents = [candidate["content"] for candidate in candidates]
        response = self._client.rerank(
            query=query,
            documents=documents,
            top_n=min(len(documents), max(top_n, 1)),
            model=self._model,
        )
        ordered: List[Dict[str, object]] = []
        for result in response.results:
            ordered.append(candidates[result.index])
        return ordered


class LocalReranker(BaseReranker):
    SECTION_PATTERN = re.compile(r"\b\w+\b")

    def rerank(
        self,
        query: str,
        candidates: Sequence[Dict[str, object]],
        top_n: int,
    ) -> List[Dict[str, object]]:
        if not candidates:
            return []
        query_terms = set(self.SECTION_PATTERN.findall(query.lower()))
        scored: List[Tuple[float, Dict[str, object]]] = []
        for candidate in candidates:
            chunk_terms = set(self.SECTION_PATTERN.findall(candidate["content"].lower()))
            overlap = len(query_terms & chunk_terms)
            score = overlap / max(len(query_terms), 1)
            scored.append((score, candidate))
        scored.sort(key=lambda item: item[0], reverse=True)
        return [item[1] for item in scored[:top_n]]


def _build_splitter(model_name: str, chunk_size: int, overlap: int, fallback_multiplier: int = 4) -> RecursiveCharacterTextSplitter:
    if chunk_size <= 0:
        chunk_size = 600
    if overlap < 0:
        overlap = 0
    try:
        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            model_name=model_name,
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n\n", "\n", " ", ""],
        )
        return splitter
    except Exception:  # pragma: no cover - fall back to character-based splitter
        return RecursiveCharacterTextSplitter(
            chunk_size=max(chunk_size * fallback_multiplier, 1000),
            chunk_overlap=max(overlap * fallback_multiplier, 200),
            separators=["\n\n", "\n", " ", ""],
        )


class RetrievalPipeline:
    def __init__(
        self,
        vector_store: VectorStoreProvider,
        keyword_index: BM25KeywordIndex,
        reranker: Optional[BaseReranker],
        enable_hybrid: bool,
        hybrid_alpha: float,
        enable_rerank: bool,
        rerank_top_n: int,
    ) -> None:
        self._vector_store = vector_store
        self._keyword_index = keyword_index
        self._reranker = reranker
        self._enable_hybrid = enable_hybrid
        self._hybrid_alpha = hybrid_alpha
        self._enable_rerank = enable_rerank and reranker is not None
        self._rerank_top_n = rerank_top_n

    def add_chunks(self, chunks: Sequence[DocumentChunk]) -> None:
        if not chunks:
            return
        self._vector_store.add_documents(chunks)
        self._keyword_index.add_chunks(chunks)

    @staticmethod
    def _normalize_scores(values: Sequence[float]) -> List[float]:
        if not values:
            return []
        minimum = min(values)
        maximum = max(values)
        if maximum - minimum < 1e-9:
            return [1.0 for _ in values]
        return [(value - minimum) / (maximum - minimum) for value in values]

    def _ensemble(
        self,
        dense_results: Sequence[Dict[str, object]],
        keyword_results: Sequence[Dict[str, object]],
        k: int,
    ) -> List[Dict[str, object]]:
        combined: Dict[Tuple[str, int], Dict[str, object]] = {}
        dense_norm = self._normalize_scores([item.get("dense_score", 0.0) for item in dense_results])
        bm25_norm = self._normalize_scores([item.get("bm25_score", 0.0) for item in keyword_results])

        for idx, candidate in enumerate(dense_results):
            metadata = candidate.get("metadata", {})
            key = (
                str(metadata.get("source", metadata.get("path", "dense"))),
                int(metadata.get("chunk_index", idx)),
            )
            combined[key] = dict(candidate)
            combined[key]["_dense_norm"] = dense_norm[idx]
            combined[key]["_bm25_norm"] = 0.0

        for idx, candidate in enumerate(keyword_results):
            metadata = candidate.get("metadata", {})
            key = (
                str(metadata.get("source", metadata.get("path", "keyword"))),
                int(metadata.get("chunk_index", idx)),
            )
            if key not in combined:
                combined[key] = dict(candidate)
                combined[key]["_dense_norm"] = 0.0
            combined[key]["_bm25_norm"] = bm25_norm[idx]

        ranking: List[Dict[str, object]] = []
        for candidate in combined.values():
            dense_score = float(candidate.get("_dense_norm", 0.0))
            keyword_score = float(candidate.get("_bm25_norm", 0.0))
            candidate["_ensemble"] = self._hybrid_alpha * dense_score + (1 - self._hybrid_alpha) * keyword_score
            ranking.append(candidate)

        ranking.sort(key=lambda item: item.get("_ensemble", 0.0), reverse=True)
        return ranking[:k]

    def retrieve(self, query: str, k: int) -> List[Dict[str, object]]:
        dense_candidates = self._vector_store.similarity_search(query, k=max(k, self._rerank_top_n))
        if not self._enable_hybrid:
            ordered = dense_candidates[:k]
        else:
            keyword_candidates = self._keyword_index.search(query, k=max(k, self._rerank_top_n))
            if not keyword_candidates:
                ordered = dense_candidates[:k]
            else:
                ordered = self._ensemble(dense_candidates, keyword_candidates, k=max(k, self._rerank_top_n))
        if self._enable_rerank and ordered:
            reranked = self._reranker.rerank(query, ordered, top_n=self._rerank_top_n)
            if reranked:
                ordered = reranked
        return ordered[:k]


class RAGService:
    """Encapsulates retrieval-augmented generation for nutrition guidance."""

    SECTION_PATTERN = re.compile(r"^#{1,6}\s+(.*)$", re.MULTILINE)

    def __init__(self) -> None:
        self.settings = get_settings()

        if not self.settings.openai_api_key:
            raise RuntimeError(
                "OPENAI_API_KEY is not set. Provide a key via environment variable or .env file.",
            )

        self._data_dir: Path = self.settings.data_dir
        self._vector_store_dir: Path = self.settings.vector_store_dir
        self._vector_store_dir.mkdir(parents=True, exist_ok=True)

        self._embeddings = OpenAIEmbeddings(
            api_key=self.settings.openai_api_key,
            model=self.settings.embedding_model,
        )
        self._llm = ChatOpenAI(
            api_key=self.settings.openai_api_key,
            model=self.settings.chat_model,
            temperature=self.settings.temperature,
            max_tokens=self.settings.max_tokens,
        )
        self._summary_llm = ChatOpenAI(
            api_key=self.settings.openai_api_key,
            model=self.settings.chat_model,
            temperature=0.0,
            max_tokens=self.settings.summary_max_tokens,
        )
        self._encoding = self._resolve_encoding(self.settings.chat_model)
        splitter_model = self.settings.embedding_model if self.settings.chunker == "token" else self.settings.chat_model
        self._splitter = _build_splitter(
            splitter_model,
            chunk_size=self.settings.chunk_size_tokens,
            overlap=self.settings.chunk_overlap_tokens,
        )

        if self.settings.vector_backend == "pgvector":
            self._vector_provider: VectorStoreProvider = PgVectorStoreProvider(self.settings.pg_dsn, self._embeddings)
        else:
            store_path = self._vector_store_dir / "nutrition_knowledge"
            self._vector_provider = FaissVectorStoreProvider(store_path, self._embeddings)

        self._bm25_index = BM25KeywordIndex()
        reranker = self._build_reranker()
        self._pipeline = RetrievalPipeline(
            vector_store=self._vector_provider,
            keyword_index=self._bm25_index,
            reranker=reranker,
            enable_hybrid=self.settings.enable_hybrid,
            hybrid_alpha=self.settings.hybrid_alpha,
            enable_rerank=self.settings.enable_rerank,
            rerank_top_n=self.settings.rerank_top_n,
        )

        self._memory_service: Optional[MemoryService] = get_memory_service()
        if self.settings.enable_memory and not self._memory_service:
            raise RuntimeError("Memory is enabled but could not initialize the memory service. Check MEMORY_DSN settings.")
        if not self.settings.enable_memory:
            self._memory_service = None

        initial_chunks = self._bootstrap_corpus()
        if not initial_chunks and self._vector_provider.is_empty():
            raise RuntimeError(
                f"No seed documents found in {self._data_dir}. Add knowledge files (txt, md, pdf, docx) to bootstrap the knowledge base.",
            )

        self._prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    "You are a registered dietitian AI that provides evidence-based, actionable nutrition advice. "
                    "Cite specific guidelines from the provided context. "
                    "If the context is insufficient, say you are unsure rather than guessing.",
                ),
                (
                    "human",
                    "Conversation memory:\n{memory}\n\nContext:\n{context}\n\nQuestion: {question}\n\nRespond with a concise answer and include a bullet list of key points.",
                ),
            ],
        )
        self._summary_prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    "You are summarizing an ongoing conversation between a user and a nutrition coach assistant. "
                    "Capture key goals, constraints, dietary preferences, and decisions without inventing information. "
                    "Keep the summary under {summary_max_tokens} tokens.",
                ),
                (
                    "human",
                    "Previous summary (may be empty):\n{previous_summary}\n\nRecent conversation transcript:\n{transcript}\n\nProvide an updated summary in plain text.",
                ),
            ],
        )

    @staticmethod
    def _resolve_encoding(model_name: str):
        try:
            return tiktoken.encoding_for_model(model_name)
        except Exception:  # pragma: no cover - fallback
            return tiktoken.get_encoding("cl100k_base")

    def _build_reranker(self) -> Optional[BaseReranker]:
        if not self.settings.enable_rerank or self.settings.rerank_provider == "none":
            return None
        if self.settings.rerank_provider == "cohere":
            if not self.settings.cohere_api_key:
                return None
            try:
                return CohereReranker(self.settings.cohere_api_key)
            except Exception:
                return None
        if self.settings.rerank_provider == "local":
            return LocalReranker()
        return None

    def _bootstrap_corpus(self) -> List[DocumentChunk]:
        chunks: List[DocumentChunk] = []
        existing_chunks = self._vector_provider.list_documents()
        if existing_chunks:
            self._bm25_index.add_chunks(existing_chunks)
            chunks.extend(existing_chunks)
            return chunks

        if self._data_dir.exists():
            for path in sorted(self._data_dir.rglob("*")):
                if path.is_dir():
                    continue
                chunks_from_path = self._load_chunks_from_path(path)
                if not chunks_from_path:
                    continue
                self._pipeline.add_chunks(chunks_from_path)
                chunks.extend(chunks_from_path)
        return chunks

    def _load_chunks_from_path(self, path: Path) -> List[DocumentChunk]:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md", ".mdx"}:
            text = path.read_text(encoding="utf-8")
            meta = {
                "source": path.name,
                "path": str(path),
                "title": path.stem,
                "source_type": suffix.lstrip("."),
            }
            return self._chunk_document(text, meta)
        if suffix == ".pdf" and self.settings.loaders.get("enable_pdf", True):
            return self._load_pdf(path)
        if suffix == ".docx" and self.settings.loaders.get("enable_docx", True):
            return self._load_docx(path)
        if suffix in {".url", ".link"} and self.settings.loaders.get("enable_web", True):
            url = path.read_text(encoding="utf-8").strip()
            return self._load_url(url, title=path.stem or url)
        return []

    def _chunk_document(self, text: str, base_metadata: Dict[str, object]) -> List[DocumentChunk]:
        normalized_text = text.strip()
        if not normalized_text:
            return []
        documents = self._splitter.create_documents([normalized_text], metadatas=[base_metadata])
        total = len(documents)
        current_section = base_metadata.get("section") or base_metadata.get("title")
        chunks: List[DocumentChunk] = []
        for idx, doc in enumerate(documents):
            metadata = dict(base_metadata)
            metadata.update(doc.metadata or {})
            header_match = self.SECTION_PATTERN.search(doc.page_content)
            if header_match:
                current_section = header_match.group(1).strip()
            if current_section:
                metadata["section"] = current_section
            metadata["chunk_index"] = idx
            metadata["num_chunks"] = total
            metadata.setdefault("title", base_metadata.get("title"))
            chunk_content = doc.page_content.strip()
            if chunk_content:
                chunks.append(DocumentChunk(content=chunk_content, metadata=metadata))
        return chunks

    def _load_pdf(self, path: Path) -> List[DocumentChunk]:
        try:
            from pypdf import PdfReader
        except ImportError:  # pragma: no cover - optional dependency
            return []
        reader = PdfReader(str(path))
        chunks: List[DocumentChunk] = []
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            meta = {
                "source": path.name,
                "path": str(path),
                "title": path.stem,
                "source_type": "pdf",
                "page": page_number,
            }
            chunks.extend(self._chunk_document(text, meta))
        return chunks

    def _load_docx(self, path: Path) -> List[DocumentChunk]:
        try:
            import docx  # type: ignore
        except ImportError:  # pragma: no cover - optional dependency
            return []
        document = docx.Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)
        meta = {
            "source": path.name,
            "path": str(path),
            "title": path.stem,
            "source_type": "docx",
        }
        return self._chunk_document(text, meta)

    def _load_url(self, url: str, title: Optional[str] = None) -> List[DocumentChunk]:
        try:
            import trafilatura
        except ImportError:  # pragma: no cover - optional dependency
            return []
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            return []
        text = trafilatura.extract(downloaded) or ""
        meta = {
            "source": title or url,
            "path": url,
            "title": title or url,
            "source_type": "url",
        }
        return self._chunk_document(text, meta)

    def ingest(self, content: str, title: str | None = None) -> int:
        """Index new text content into the vector store."""
        metadata = {
            "source": title or f"upload-{uuid4().hex[:8]}",
            "title": title or "User Upload",
            "source_type": "ingest",
        }
        chunks = self._chunk_document(content, metadata)
        if not chunks:
            return 0
        self._pipeline.add_chunks(chunks)
        self._vector_provider.persist()
        return len(chunks)

    def ingest_url(self, url: str, title: Optional[str] = None) -> int:
        chunks = self._load_url(url, title=title)
        if not chunks:
            return 0
        self._pipeline.add_chunks(chunks)
        self._vector_provider.persist()
        return len(chunks)

    def answer(
        self,
        question: str,
        top_k: int | None = None,
        session_id: Optional[str] = None,
        user_id: Optional[str] = None,
    ) -> dict:
        """Run retrieval + generation and return the answer with supporting chunks."""
        k = top_k or self.settings.top_k
        if k <= 0:
            k = self.settings.top_k

        memory_text = ""
        memory_tokens = 0
        conversation_id: Optional[str] = None
        if session_id and self._memory_service:
            conversation_id = session_id
            self._memory_service.ensure_conversation(conversation_id, user_id=user_id)
            memory_text = self._build_memory_text(conversation_id)
            if memory_text:
                memory_tokens = len(self._encoding.encode(memory_text))

        retrieved = self._pipeline.retrieve(question, k)
        if not retrieved:
            answer_text = (
                "I could not find relevant information in the knowledge base yet. "
                "Please add more nutrition documents and try again."
            )
            if conversation_id and self._memory_service:
                self._persist_turn(conversation_id, question, answer_text)
            return {"answer": answer_text, "context": []}

        available_tokens = self.settings.max_input_tokens - memory_tokens
        if available_tokens <= 0:
            available_tokens = self.settings.chunk_size_tokens

        selected_chunks, context = self._build_context(retrieved, available_tokens)
        prompt_memory = memory_text or "None"
        messages = self._prompt.format_messages(memory=prompt_memory, context=context, question=question)
        response = self._llm.invoke(messages)
        answer_text = response.content

        if conversation_id and self._memory_service:
            self._persist_turn(conversation_id, question, answer_text)

        return {
            "answer": answer_text,
            "context": selected_chunks,
        }

    def _build_context(self, chunks: Sequence[Dict[str, object]], token_budget: int) -> Tuple[List[Dict[str, object]], str]:
        selected: List[Dict[str, object]] = []
        token_budget = max(token_budget, 0)
        tokens_used = 0
        for chunk in chunks:
            content = chunk.get("content", "")
            if not isinstance(content, str) or not content.strip():
                continue
            token_count = len(self._encoding.encode(content))
            if selected and tokens_used + token_count > token_budget:
                break
            tokens_used += token_count
            metadata = dict(chunk.get("metadata", {}))
            selected.append({"content": content, "metadata": metadata})
        if not selected and chunks:
            first = chunks[0]
            selected.append(
                {
                    "content": first.get("content", ""),
                    "metadata": dict(first.get("metadata", {})),
                },
            )
        context = "\n\n".join(entry["content"] for entry in selected)
        return selected, context

    def _build_memory_text(self, conversation_id: str) -> str:
        if not self._memory_service:
            return ""
        try:
            summary = self._memory_service.get_summary(conversation_id)
            history_limit = max(self.settings.history_window * 2, 0)
            recent_messages = self._memory_service.get_recent_messages(conversation_id, history_limit)
        except Exception:  # pragma: no cover - memory fetch failure should not break answer
            return ""

        sections: List[str] = []
        if summary:
            sections.append(f"Summary:\n{summary.strip()}")
        if recent_messages:
            dialog_lines = []
            for message in recent_messages:
                speaker = "User" if message.role == "user" else "Assistant"
                dialog_lines.append(f"{speaker}: {message.content.strip()}")
            sections.append("Recent dialogue:\n" + "\n".join(dialog_lines))
        return "\n\n".join(sections)

    def _persist_turn(self, conversation_id: str, user_message: str, assistant_message: str) -> None:
        if not self._memory_service:
            return
        try:
            self._memory_service.append_message(conversation_id, "user", user_message)
            self._memory_service.append_message(conversation_id, "assistant", assistant_message)
            self._memory_service.set_title_if_absent(conversation_id, user_message[:80])
            self._maybe_update_summary(conversation_id)
        except Exception:  # pragma: no cover - persistence failures should not crash
            return

    def _maybe_update_summary(self, conversation_id: str) -> None:
        if not self._memory_service:
            return
        try:
            total_messages = self._memory_service.get_message_count(conversation_id)
        except Exception:  # pragma: no cover
            return
        turns = total_messages // 2
        if turns == 0 or turns % self.settings.summary_every_n_turns != 0:
            return

        try:
            existing_summary = self._memory_service.get_summary(conversation_id) or ""
            messages = self._memory_service.iter_messages(conversation_id)
            window = max(self.settings.history_window * 4, 20)
            if window > 0:
                messages = messages[-window:]
            transcript_lines = []
            for message in messages:
                speaker = "User" if message.role == "user" else "Assistant"
                transcript_lines.append(f"{speaker}: {message.content.strip()}")
            transcript = "\n".join(transcript_lines)
            summary_messages = self._summary_prompt.format_messages(
                previous_summary=existing_summary or "(none)",
                transcript=transcript,
                summary_max_tokens=self.settings.summary_max_tokens,
            )
            result = self._summary_llm.invoke(summary_messages)
            summary_text = result.content.strip()
            self._memory_service.set_summary(conversation_id, summary_text)
        except Exception:  # pragma: no cover
            return



def get_rag_service() -> RAGService:
    # Lazy-load so FastAPI dependency injection can manage a single instance
    if not hasattr(get_rag_service, "_instance"):
        get_rag_service._instance = RAGService()
    return get_rag_service._instance  # type: ignore[attr-defined]
